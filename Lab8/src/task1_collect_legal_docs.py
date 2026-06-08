"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Dùng Crawl4AI để crawl toàn văn từ các nguồn chính thống (vbpl.vn, thuvienphapluat.vn).
Nếu trang có link PDF/DOC gốc thì tải về; nếu không thì lưu nội dung crawl thành .docx.
"""

import asyncio
import re
from pathlib import Path
from urllib.parse import urljoin

import requests

from src.crawl_utils import DEFAULT_CRAWL_CONFIG

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"

# Nguon chinh thong - toan van cong khai tren thuvienphapluat.vn
LEGAL_SOURCES = [
    {
        "url": "https://thuvienphapluat.vn/van-ban/Trach-nhiem-hinh-su/Luat-Phong-chong-ma-tuy-2021-445185.aspx",
        "filename": "luat-phong-chong-ma-tuy-2021",
        "title": "Luat 73/2021/QH14 - Phong chong ma tuy",
    },
    {
        "url": "https://thuvienphapluat.vn/van-ban/Van-hoa-Xa-hoi/Nghi-dinh-105-2021-ND-CP-huong-dan-Luat-Phong-chong-ma-tuy-496664.aspx",
        "filename": "nghi-dinh-105-2021",
        "title": "Nghi dinh 105/2021/ND-CP",
    },
    {
        "url": "https://thuvienphapluat.vn/van-ban/Trach-nhiem-hinh-su/Bo-luat-Hinh-su-sua-doi-bo-sung-2017-359220.aspx",
        "filename": "bo-luat-hinh-su-2015",
        "title": "Bo luat Hinh su 2015 (sua doi 2017)",
    },
    {
        "url": "https://thuvienphapluat.vn/van-ban/Van-hoa-Xa-hoi/Nghi-dinh-57-2022-ND-CP-danh-muc-chat-ma-tuy-va-tien-chat-527507.aspx",
        "filename": "nghi-dinh-57-2022-danh-muc-ma-tuy",
        "title": "Nghi dinh 57/2022/ND-CP - Danh muc chat ma tuy",
    },
    {
        "url": "https://thuvienphapluat.vn/van-ban/Van-hoa-Xa-hoi/Nghi-dinh-90-2024-ND-CP-sua-doi-Danh-muc-chat-ma-tuy-tien-chat-theo-Nghi-dinh-57-2022-ND-CP-607161.aspx",
        "filename": "nghi-dinh-90-2024-sua-doi-danh-muc",
        "title": "Nghi dinh 90/2024/ND-CP sua doi danh muc ma tuy",
    },
]


def setup_directory() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[OK] Thu muc: {DATA_DIR}")


def _download_binary(url: str, filepath: Path) -> bool:
    """Try downloading a PDF/DOC file from a direct link."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        )
    }
    try:
        response = requests.get(url, headers=headers, timeout=90)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "")
        if len(response.content) < 1024:
            return False
        if "html" in content_type.lower() and not url.lower().endswith((".pdf", ".doc", ".docx")):
            return False
        filepath.write_bytes(response.content)
        print(f"[OK] Da tai file goc: {filepath} ({len(response.content):,} bytes)")
        return True
    except requests.RequestException as exc:
        print(f"[WARN] Khong tai duoc {url}: {exc}")
        return False


def _save_as_docx(content: str, filepath: Path, title: str) -> Path:
    """Save crawled markdown/text as DOCX."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in content.split("\n"):
        text = paragraph.strip()
        if text:
            doc.add_paragraph(text)
    doc.save(filepath)
    print(f"[OK] Luu noi dung crawl thanh DOCX: {filepath} ({filepath.stat().st_size:,} bytes)")
    return filepath


def _extract_download_links(result, base_url: str) -> list[str]:
    """Find PDF/DOC download links from crawled page."""
    links: list[str] = []
    patterns = (r"\.pdf", r"\.docx?", r"ViewFile\.aspx", r"download", r"van-ban-goc")

    candidates = []
    if result.links:
        for category in ("internal", "external"):
            for item in result.links.get(category, []):
                href = item.get("href") if isinstance(item, dict) else str(item)
                if href:
                    candidates.append(href)

    html = result.html or ""
    candidates.extend(re.findall(r'href=["\']([^"\']+)["\']', html, flags=re.I))

    for href in candidates:
        full_url = urljoin(base_url, href)
        lower = full_url.lower()
        if any(p in lower for p in patterns):
            links.append(full_url)

    # Deduplicate while preserving order
    seen = set()
    unique = []
    for link in links:
        if link not in seen:
            seen.add(link)
            unique.append(link)
    return unique


async def crawl_legal_document(source: dict) -> Path:
    """Crawl one legal document page with Crawl4AI."""
    from crawl4ai import AsyncWebCrawler

    url = source["url"]
    base_name = source["filename"]

    async with AsyncWebCrawler(verbose=False) as crawler:
        result = await crawler.arun(url=url, config=DEFAULT_CRAWL_CONFIG)

    if not result.success:
        raise RuntimeError(f"Crawl that bai: {url} - {result.error_message}")

    content = result.markdown or result.cleaned_html or ""
    if len(content.strip()) < 500:
        raise RuntimeError(f"Noi dung crawl qua ngan ({len(content)} chars): {url}")

    # Thu tai PDF/DOC goc neu co link tren trang
    for link in _extract_download_links(result, url):
        ext = ".pdf"
        if link.lower().endswith(".docx"):
            ext = ".docx"
        elif link.lower().endswith(".doc"):
            ext = ".doc"
        filepath = DATA_DIR / f"{base_name}{ext}"
        if _download_binary(link, filepath):
            return filepath

    # Khong co file goc -> luu noi dung crawl thanh DOCX (du lieu that tu nguon chinh thong)
    docx_path = DATA_DIR / f"{base_name}.docx"
    return _save_as_docx(content, docx_path, source["title"])


async def collect_all(skip_existing: bool = True) -> list[Path]:
    """Crawl all legal documents."""
    setup_directory()
    collected: list[Path] = []

    for i, source in enumerate(LEGAL_SOURCES, 1):
        existing = list(DATA_DIR.glob(f"{source['filename']}.*"))
        if skip_existing and existing:
            valid = [p for p in existing if p.stat().st_size > 1024]
            if valid:
                print(f"[{i}/{len(LEGAL_SOURCES)}] [SKIP] Da co: {valid[0].name}")
                collected.append(valid[0])
                continue

        print(f"[{i}/{len(LEGAL_SOURCES)}] Crawling: {source['url']}")
        path = await crawl_legal_document(source)
        collected.append(path)

    return collected


def download_all(skip_existing: bool = True) -> list[Path]:
    """Sync entry point (giu ten ham cu de tuong thich)."""
    return asyncio.run(collect_all(skip_existing=skip_existing))


if __name__ == "__main__":
    files = download_all(skip_existing=False)
    print(f"\nTong cong: {len(files)} file trong {DATA_DIR}")
    for f in files:
        print(f"  - {f.name} ({f.stat().st_size:,} bytes)")
